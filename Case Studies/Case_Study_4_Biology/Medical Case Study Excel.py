import math
import pandas as pd
from pandas import ExcelWriter
from pandas import ExcelFile
import docx

# Data = {
#     "symptoms": ["fever", "chills", "headache", "diarrhea", "constipation", "anorexia", "malaise", "cough", "vomiting",
#                  "abdominal pain", "Hepatomegaly", "Gastrointestial bleeding", "Changes in sensorium", "Rashes",
#                  "seizures"],
#     "blood culture": {
#         "sensitivity": [17.3, 49.23],
#         "specificity": [91.1, 100],
#         "ppv": [70.1, 100],
#         "npv": [50.4, 75.1]
#     },
#     "widal test": {
#         "sensitivity": [25.5, 59.26],
#         "specificity": [77.6, 97.0],
#         "ppv": [54.8, 92.9],
#         "npv": [51.3, 77.1]
#     },
#     "typhiod M test": {
#         "sensitivity": [69.4, 94.5],
#         "specificity": [90.1, 100],
#         "ppv": [86.7, 100],
#         "npv": [76.4, 95.9]
#     },
#     "Diazo test": {
#         "sensitivity": [61.6, 90.2],
#         "specificity": [70.6, 93.7],
#         "ppv": [64.4, 92.1],
#         "npv": [68.1, 92.1]
#     }
# }
symptom_dic = {}
standardvalues = {}
uservalues = {}


def calcSymptomPercentage(count):
    return (count / 15) * 100


def calcTestPercentage(testcount):
    return (testcount / 4) * 100


# to get double value of dictinary
# print(Data["blood culture"]["sensitivity"])

def read_symptom_data(df):
    symlist = df['Symptoms']
    newlist = list()
    # print(symlist)

    for i in range(1, symlist.__len__() + 1):
        newlist.append(symlist[i])

    symptom_dic['Symptoms'] = newlist
    # print(symptom_dic)


def read_test_values(df):
    senstivitylist = df['SENSITIVITY']
    speicivitylist = df['SPECIFICITY']
    ppvlist = df['PPV']
    npvlist = df['NPV']

    standardvalues["Blood culture"] = {}
    standardvalues["widal test"] = {}
    standardvalues["typhiod M test"] = {}
    standardvalues["Diazo test"] = {}
    startval = 1
    for key in standardvalues.keys():

        tempsens = list()
        tempspef = list()
        tempppv = list()
        tempnpv = list()

        for x in range(startval, (startval + 2)):
            tempsens.append(senstivitylist[x])
            tempspef.append(speicivitylist[x])
            tempppv.append(ppvlist[x])
            tempnpv.append(npvlist[x])

            # print(x)

        startval = startval + 2
        standardvalues[key]['sensitivity'] = tempsens
        standardvalues[key]['specificity'] = tempspef
        standardvalues[key]['ppv'] = tempnpv
        standardvalues[key]['npv'] = tempnpv

    # print(standardvalues)
    # for i in range(1,13,2):
    #     print(i)
    #     i=i+1
    #     print(i)


def read_user_test_values(df):
    senstivitylist = df['SENSITIVITY']
    speicivitylist = df['SPECIFICITY']
    ppvlist = df['PPV']
    npvlist = df['NPV']

    uservalues["Blood culture"] = {}
    uservalues["widal test"] = {}
    uservalues["typhiod M test"] = {}
    uservalues["Diazo test"] = {}
    startval = 10
    for key in uservalues.keys():

        tempsens = list()
        tempspef = list()
        tempppv = list()
        tempnpv = list()

        for x in range(startval, (startval + 1)):
            tempsens.append(senstivitylist[x])
            tempspef.append(speicivitylist[x])
            tempppv.append(ppvlist[x])
            tempnpv.append(npvlist[x])

            # print(x)

        startval = startval + 2
        uservalues[key]['sensitivity'] = tempsens
        uservalues[key]['specificity'] = tempspef
        uservalues[key]['ppv'] = tempnpv
        uservalues[key]['npv'] = tempnpv

    # print(uservalues)
    # for i in range(1,13,2):
    #     print(i)
    #     i=i+1
    #     print(i)


def read_excel_values():
    dataframe = pd.read_excel('NormalValues.xlsx')
    # dataframe=dataframe.dropna(axis=0, how='any')
    # print(dataframe['Test Name'])
    # print(dataframe.isnull().sum())
    # dataframe.dropna(axis=0,subset=['Test Name'],how='any')
    # print(dataframe[dataframe['Test Name'].notnull])
    symptomdf = dataframe[dataframe['Symptoms'].notnull()]
    read_symptom_data(symptomdf)
    read_test_values(dataframe)
    read_user_test_values(dataframe)


def getinput():
    read_excel_values()
    print("Symptoms for typhoid are given below\nEnter  yes or 0 for no")
    #     # Input symptoms
    count = 0
    symptom_value_list = list()

    for x in symptom_dic['Symptoms']:
        x = (input(x + ": "))
        symptom_value_list.append(x)
        if x.lower() == 'yes':
            count = count + 1

    symptom_dic['Symptom Values'] = symptom_value_list
    symptom_dic['Count'] = count
    # print(symptom_dic)


#     user_dic = dict()
#     # Input test results
#     for x in Data:
#         if x == "symptoms":
#             continue
#         #     to skip symptoms one
#         else:
#             print("Enter " + x + " values")
#             user_dic[x] = (input("sensitivity value"), input("specificity"), input("ppv"), input("npv"))
#
#     return (count, user_dic)
#
#
def compute(symptomCount, user_dic):
    # print("Count is {}".format(symptomCount))
    symptom_dic['Percentage'] = calcSymptomPercentage(symptomCount)

    # Compute test results
    for x in standardvalues:

        positive_test_count = 0
        for y in standardvalues[x]:

            lowerBound = float(standardvalues[x][y][0])
            upperBound = float(standardvalues[x][y][1])
            userInput = float(user_dic[x][y][0])

            if userInput >= lowerBound and userInput <= upperBound:
                positive_test_count = positive_test_count + 1

        print("test count {}".format(positive_test_count))
        # testResult[x] = calcTestPercentage(positive_test_count)
        testres=list()
        testres.append(calcTestPercentage(positive_test_count))
        user_dic[x]['Test Result'] = testres


#
#
def output():
    doc = docx.Document()
    doc.add_heading('Medical Case Study Report', 0)
    doc.add_heading('Symptom Results', 1)
    table = doc.add_table(1, 2)
    table.style = 'MediumGrid3-Accent2'
    table.cell(0, 0).text = 'Symptoms'
    table.cell(0, 1).text = 'Results'

    #add symptom and result data
    for x in range(len(symptom_dic['Symptoms']) ):
        sympval=symptom_dic['Symptoms'][x]
        answerval=symptom_dic['Symptom Values'][x]
        cells = table.add_row().cells
        cells[0].text=sympval
        cells[1].text = answerval

    cells = table.add_row().cells
    cells[0].text = 'No of symptoms found'
    cells[1].text = format(symptom_dic['Count'])




    cells = table.add_row().cells
    cells[0].text = 'Percentage'
    cells[1].text = format(symptom_dic['Percentage'],'.2f')+"%"

    doc.add_page_break()

    doc.add_heading('Test Results', 1)
    table = doc.add_table(1, 6)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Diagnostic Tests'
    table.cell(0, 1).text = 'Sensitivity'
    table.cell(0, 2).text = 'Specifivity'
    table.cell(0, 3).text = 'PPV'
    table.cell(0, 4).text = 'NPV'
    table.cell(0, 5).text = 'Percentage found'

    # print user data
    for x in uservalues:
        cells = table.add_row().cells
        cells[0].text = x
        startval=1
        for y in uservalues[x]:
            cells[startval].text=format(uservalues[x][y][0])
            startval=startval+1


    doc.save('testfile.docx')

    print(symptom_dic)
    print(standardvalues)
    print(uservalues)


#
#
# C, D = getinput()
# P, R = compute(C, D)
# output(P, R)

getinput()
compute(symptom_dic['Count'], uservalues)
output()

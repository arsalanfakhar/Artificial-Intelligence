import docx

doc=docx.Document()

doc.add_heading('Heading for the document', 0)

p = doc.add_paragraph('A plain˓paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

#make a table
table=doc.add_table(rows=3,cols=2)

for x in range(len(table.rows)):
    for y in range(len(table.columns)):
        table.cell(x,y).text='kuchi'


table.style ='Table Grid'
doc.save('testfile.docx')

